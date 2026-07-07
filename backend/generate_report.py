"""
将校园二手交易平台实训报告导出为 .docx
基于 python-docx 1.2.0
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ============= 样式工具 =============

def set_cell_border(cell, **kwargs):
    """给单元格设置边框。kwargs: top/left/bottom/right -> dict(sz, color)"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = tcPr.find(qn('w:tcBorders'))
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)
    for edge in ('top', 'left', 'bottom', 'right'):
        if edge in kwargs:
            border = tcBorders.find(qn(f'w:{edge}'))
            if border is None:
                border = OxmlElement(f'w:{edge}')
                tcBorders.append(border)
            for k, v in kwargs[edge].items():
                border.set(qn(f'w:{k}'), str(v))


def set_cell_shading(cell, fill_hex):
    """设置单元格背景色。"""
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_hex)
    tcPr.append(shd)


def add_para(doc, text, *, size=10.5, bold=False, italic=False,
             align=None, first_line_indent=True, space_after=4,
             color=None, line_spacing=1.5):
    """添加普通段落。"""
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    pf = p.paragraph_format
    pf.line_spacing = line_spacing
    pf.space_after = Pt(space_after)
    if first_line_indent:
        pf.first_line_indent = Cm(0.74)  # 2 字符
    run = p.add_run(text)
    run.font.name = '宋体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color
    return p


def add_heading(doc, text, level=1):
    """添加标题。level: 1=章 2=节 3=小节"""
    p = doc.add_paragraph()
    if level == 1:
        # 一级标题：居中，黑体，16pt
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        size = 16
        run = p.add_run(text)
        run.font.name = '黑体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        run.bold = True
    elif level == 2:
        # 二级标题：左对齐，黑体，14pt
        size = 14
        run = p.add_run(text)
        run.font.name = '黑体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        run.bold = True
    else:
        # 三级标题：左对齐，黑体，12pt
        size = 12
        run = p.add_run(text)
        run.font.name = '黑体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        run.bold = True
    run.font.size = Pt(size)
    pf = p.paragraph_format
    pf.space_before = Pt(10)
    pf.space_after = Pt(6)
    pf.line_spacing = 1.5
    return p


def add_code_block(doc, code_text, lang=''):
    """添加代码块。"""
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.left_indent = Cm(0.5)
    pf.space_after = Pt(4)
    pf.line_spacing = 1.15
    # 浅灰底
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'F5F5F5')
    pPr.append(shd)
    lines = code_text.split('\n')
    for i, line in enumerate(lines):
        if i > 0:
            r = p.add_run()
            r.add_break()
        run = p.add_run(line if line else ' ')
        run.font.name = 'Consolas'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Consolas')
        run.font.size = Pt(9)
    return p


def add_table(doc, header, rows, col_widths=None):
    """添加表格，header 列表，rows 是 list[list[str]]。"""
    table = doc.add_table(rows=1 + len(rows), cols=len(header))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'
    # 表头
    for i, h in enumerate(header):
        cell = table.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        run.bold = True
        run.font.name = '黑体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        run.font.size = Pt(10.5)
        set_cell_shading(cell, 'D9E2F3')
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    # 数据
    for r_idx, row in enumerate(rows, start=1):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx].cells[c_idx]
            cell.text = ''
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(str(val))
            run.font.name = '宋体'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            run.font.size = Pt(10)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Cm(w)
    return table


def add_page_break(doc):
    doc.add_page_break()


# ============= 文档主体 =============

doc = Document()

# 默认样式：A4 + 宋体小四
for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

# Normal 样式
style = doc.styles['Normal']
style.font.name = '宋体'
style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
style.font.size = Pt(10.5)

# ============= 封面 =============
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
pf = p.paragraph_format
pf.space_before = Pt(120)
run = p.add_run('校园二手交易平台')
run.font.name = '黑体'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
run.font.size = Pt(28)
run.bold = True
run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
pf = p.paragraph_format
pf.space_before = Pt(12)
run = p.add_run('— 课设项目实训报告 —')
run.font.name = '黑体'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
run.font.size = Pt(16)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
pf = p.paragraph_format
pf.space_before = Pt(60)
run = p.add_run('综合实训项目组')
run.font.size = Pt(14)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
pf = p.paragraph_format
pf.space_before = Pt(12)
run = p.add_run('2026 年 6 月')
run.font.size = Pt(12)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
pf = p.paragraph_format
pf.space_before = Pt(120)
run = p.add_run('技术栈：Django 5 + DRF + SimpleJWT + MySQL 9.4\nVue 3 + TypeScript + Element Plus + ECharts\n微信原生小程序（自定义 5 tab）')
run.font.size = Pt(11)
run.font.color.rgb = RGBColor(0x59, 0x59, 0x59)

add_page_break(doc)

# ============= 目录占位 =============
add_heading(doc, '目  录', level=1)
toc_items = [
    '第1章 项目概述',
    '第2章 系统总体设计',
    '第3章 数据库设计',
    '第4章 系统环境搭建',
    '第5章 核心模块设计与实现',
    '第6章 系统测试与三端联调',
    '第7章 开发问题与解决方案',
    '第8章 项目总结',
]
for item in toc_items:
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.line_spacing = 1.8
    pf.space_after = Pt(2)
    run = p.add_run(item)
    run.font.size = Pt(12)
    run.font.name = '宋体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
add_page_break(doc)

# ============= 第1章 =============
add_heading(doc, '第1章 项目概述', level=1)

add_heading(doc, '1.1 项目背景', level=2)
add_para(doc, '随着高校扩招与互联网消费习惯的普及，校园内闲置教材、电子产品、生活用品的二手流转需求快速增长，但目前仍以「宿舍楼贴海报」「QQ 群接龙」「朋友圈吆喝」为主，普遍存在以下痛点：')
add_para(doc, '1. 信息散落、检索困难：商品信息发布在多个群里，缺乏统一检索与类目化展示，潜在买家无法高效比价。')
add_para(doc, '2. 信任成本高：买卖双方多为陌生同学，缺乏身份核验、信用沉淀与售后保障机制，纠纷难以追溯。')
add_para(doc, '3. 交易流程不规范：从「想要购买」到「线下自取」之间缺少状态管理，经常出现「卖家已答应、买家却放鸽子」「商品被多人同时锁定」等问题。')
add_para(doc, '4. 运营监管缺位：违规信息（违禁品、虚假宣传、价格异常）缺少审核与举报通道，平台方无法介入。')
add_para(doc, '5. 教学场景契合度低：传统课程实训的「家庭记账」类项目偏个人工具，缺少复杂状态机、权限体系与多端协同，训练价值有限。')
add_para(doc, '为解决上述问题，并结合「微信小程序 + Vue3 + Django」课程教学目标，本项目将原「家庭记账」业务整体转型为 C2C 校园二手交易平台，构建一个支持买家、卖家、平台管理员三类角色，覆盖「发布 → 审核 → 私聊 → 下单 → 评价 → 信用沉淀」完整业务闭环的多端系统。')

add_heading(doc, '1.2 项目目标', level=2)
add_para(doc, '本项目旨在完成以下核心目标：')
add_para(doc, '1. 业务目标：搭建一个真实的校园二手交易平台，提供商品发布 / 浏览 / 私聊 / 下单 / 评价 / 信用分 / 平台审核的全链路功能。')
add_para(doc, '2. 技术目标：')
add_para(doc, '（1）后端使用 Django 5 + DRF + SimpleJWT + MySQL，实现 RESTful API、JWT 鉴权、对象级权限、ORM 聚合统计、事务一致性等工程实践。')
add_para(doc, '（2）前端实现「一后端 + 三前端」：微信小程序（学生端）、Vue3 + Element Plus（卖家工作台）、Vue3 + Element Plus（平台管理后台）。')
add_para(doc, '（3）集成 AI 一键发布（图片识物 + 描述润色 + 建议价）与 AI 议价参考，并实现无 Key 时的 mock 降级。')
add_para(doc, '3. 教学目标：覆盖课程 4 次实训的全部技术点（数据库设计、JWT 鉴权、状态机、聚合统计、跨端联调），并通过统一的设计 Token 体系完成 UI 标准化训练。')

add_heading(doc, '1.3 目标用户', level=2)
add_para(doc, '本项目区分三类用户角色，每类角色的需求与使用端对应关系如下：')
add_table(doc,
    ['角色', '角色标识', '典型用户', '核心需求', '使用端'],
    [
        ['普通买家', 'user', '有购买需求的学生', '浏览商品、收藏、私聊、议价、下单、评价', '微信小程序'],
        ['普通卖家', 'user', '想处理闲置物品的学生', '发布商品、管理上下架、回复私聊、确认订单、查看销售看板', '微信小程序 + Web 卖家工作台'],
        ['平台管理员', 'admin', '学校网管 / 平台运营', '审核商品、处理举报、封禁用户、管理分类、查看平台看板、配置 AI', 'Web 管理后台'],
    ],
    col_widths=[2, 2, 3, 5, 3]
)
add_para(doc, '同一自然人在不同场景下可同时是买家与卖家，因此 role 字段为单值枚举（user / admin），买家与卖家身份通过 Order.buyer 与 Order.seller 在业务层动态判定。')

add_heading(doc, '1.4 功能范围', level=2)
add_para(doc, '按照课程「必做 + 选做」要求，本项目功能范围划分如下：')
add_para(doc, '必做功能（核心 5 个模块）：', bold=True)
add_para(doc, '1. 用户与认证：注册 / 登录 / JWT 鉴权 / 个人资料维护 / 校园身份认证 / 信用分展示。')
add_para(doc, '2. 商品管理：商品发布 / 状态机（草稿、待审核、在售、已订、已售、下架）/ 类目筛选 / 关键词搜索 / 收藏 / 浏览数。')
add_para(doc, '3. 私聊与议价：基于商品的会话、消息发送、默认招呼语、未读数维护。')
add_para(doc, '4. 订单管理：下单、卖家确认、自取 / 快递、完成 / 取消、订单状态机可视化。')
add_para(doc, '5. 消费 / 销售统计：卖家看板（ECharts 销售趋势、类目分布、价格区间）、平台看板（用户数、商品数、订单数、今日新增）。')
add_para(doc, '选做功能（提升项目亮点）：', bold=True)
add_para(doc, '1. AI 一键发布：用户上传商品图 → 后端调用 LLM → 返回类目 / 标题 / 描述 / 建议价 / 置信度，前端一键回填。')
add_para(doc, '2. AI 议价参考 / AI 内容审核 / AI 客服：私聊与发布场景下的 AI 辅助。')
add_para(doc, '3. Web 平台管理后台：商品审核、举报处理、用户封禁、分类管理、审计日志、AI 配置。')
add_para(doc, '4. 信用分体系：借鉴芝麻信用，初始 80，好评 +1、差评 -1，徽章 + 数字滚动可视化。')
add_para(doc, '5. 图片上传与裁剪：商品最多 9 张图，封面、排序可在前端编辑。')

add_heading(doc, '1.5 项目特色', level=2)
add_para(doc, '本项目区别于「普通记账软件」或「普通电商模板」的两大核心亮点：')
add_para(doc, '特色 1：AI 一键发布 + 信用分体系（产品力差异）', bold=True)
add_para(doc, '借鉴闲鱼 / 得物级别的「懒人发布」体验：用户上传 1 张商品图，后端调用 LLM 自动识别物品类目、润色描述、给出建议价，前端一键回填到发布表单；同时引入「芝麻信用」式的动态信用分机制，订单完成后根据评价自动 +1 / -1，并通过彩色徽章 + 数字滚动动效可视化，让陌生人交易有据可依。')
add_para(doc, '特色 2：一后端 + 三前端的端到端联调（工程力差异）', bold=True)
add_para(doc, '同一套 Django REST API 同时服务「微信小程序（学生买卖端）+ Vue3 卖家工作台 + Vue3 平台管理后台」三类用户角色，覆盖买家、卖家、管理员全链路；通过统一的设计 Token（CSS 变量 + SVG 图标）实现三端视觉一致，并通过 PowerShell 一键脚本（setup_database.ps1 / start_all.ps1）实现「30 秒环境就绪、5 分钟三端联调演示」。')

add_page_break(doc)

# ============= 第2章 =============
add_heading(doc, '第2章 系统总体设计', level=1)

add_heading(doc, '2.1 功能架构设计', level=2)
add_heading(doc, '2.1.1 模块划分总览', level=3)
add_table(doc,
    ['一级模块', '子功能', '用户角色'],
    [
        ['鉴权模块', '注册 / 登录 / 注销 / Token 刷新 / 密码修改', '全部'],
        ['用户中心', '个人资料 / 头像上传 / 校园认证 / 信用分展示 / 我的统计', '普通用户'],
        ['分类管理', '一级 + 二级类目树 / 排序 / 启用', '普通用户（读）/ 管理员（写）'],
        ['商品管理', '发布 / 编辑 / 上下架 / 浏览数 / 收藏 / 搜索 / 状态机', '普通用户'],
        ['私聊中心', '会话列表 / 消息发送 / 未读数 / 标记已读', '普通用户'],
        ['订单中心', '下单 / 卖家确认 / 完成 / 取消 / 评价', '普通用户'],
        ['数据统计', '卖家看板（趋势 / 类目 / 价格区间）/ 平台看板', '普通用户 / 管理员'],
        ['举报中心', '创建举报 / 状态机 / 处理动作（警告 / 下架 / 封禁）', '普通用户 / 管理员'],
        ['平台管理', '商品审核 / 用户封禁 / 信用分调整 / 审计日志', '管理员'],
        ['AI 能力', '一键发布 / 议价 / 审核 / 客服 / Key 配置', '全部'],
        ['系统级', '轮播图 / 公告 / 热门关键词 / 首页 Feed / 站点统计', '全部'],
    ],
    col_widths=[2.5, 8, 4]
)

add_heading(doc, '2.1.2 思维导图（文字版）', level=3)
add_code_block(doc, '''校园二手交易平台
├── 买家侧（微信小程序）
│   ├── 首页（瀑布流 / 搜索 / 公告 / 轮播 / 热门关键词）
│   ├── 分类（一级 + 二级 tab）
│   ├── 发布（AI 一键发布 / 手动填写 / 图片上传）
│   ├── 消息（会话列表 / 聊天 / 议价 / 订单卡片）
│   └── 我的（个人资料 / 信用分徽章 / 我的发布 / 我的收藏 / 我的订单 / 设置）
├── 卖家侧（Web 卖家工作台，端口 3000）
│   ├── 卖家看板（销售趋势 / 类目分布 / 价格区间）
│   ├── 我的商品（CRUD / 上下架 / 批量操作）
│   ├── 我的订单（卖家视角 / 确认 / 完成）
│   └── 个人中心
└── 平台侧（Web 管理后台，端口 5173）
    ├── 平台 Dashboard
    ├── 用户管理（封禁 / 解封 / 调分）
    ├── 商品审核（待审核列表 / 通过 / 驳回 / 理由模板）
    ├── 分类管理
    ├── 举报处理
    ├── 审计日志
    └── AI 配置（Key 接入 / 健康检查 / Mock 开关）''')

add_heading(doc, '2.2 技术架构设计', level=2)
add_heading(doc, '2.2.1 三层架构图', level=3)
add_code_block(doc, '''┌──────────────────────────────────────────────────────────┐
│                       表现层（前端）                       │
│  ┌──────────────┐  ┌──────────────┐  ┌──────────────┐     │
│  │ 微信小程序     │  │ Vue3 卖家台   │  │ Vue3 管理后台 │     │
│  │ 端口: 微信IDE │  │ 端口 3000     │  │ 端口 5173     │     │
│  └──────┬───────┘  └──────┬───────┘  └──────┬───────┘     │
│         │ HTTPS / wx.request              │               │
└─────────┼─────────────────┼────────────────┼─────────────┘
          ▼                 ▼                ▼
┌──────────────────────────────────────────────────────────┐
│                  业务层（Django + DRF + JWT）              │
│  ├─ market/views/   (auth / product / order / ...)        │
│  ├─ market/permissions/   (IsOwner / IsAdmin)             │
│  └─ market/services/      (AI / LLM / ASR 客户端)         │
│  waitress · 监听 0.0.0.0:8000                              │
└─────────────────────────┬────────────────────────────────┘
                          │ PyMySQL / Django ORM
                          ▼
┌──────────────────────────────────────────────────────────┐
│   MySQL 9.4（campus_market） · MEDIA_ROOT · RabbitMQ      │
└──────────────────────────────────────────────────────────┘''')

add_heading(doc, '2.2.2 各层技术栈与版本', level=3)
add_table(doc,
    ['层级', '技术', '版本', '说明'],
    [
        ['前端-小程序', '微信原生 WXML / WXSS / JS', '基础库 3.x', '5 tab 自定义 tab-bar'],
        ['前端-卖家台', 'Vue3 / Vite / TS / Element Plus / ECharts / Pinia', 'Vue 3.5', '端口 3000'],
        ['前端-管理后台', 'Vue3 / Vite / JS / Element Plus / Vue Router / Pinia', 'Vue 3.5', '端口 5173'],
        ['后端-Web', 'Django', '5.x', 'MVT 架构'],
        ['后端-API', 'Django REST Framework', '13.x+', 'ViewSet + Router'],
        ['后端-鉴权', 'djangorestframework-simplejwt', '5.x', 'Access + Refresh Token'],
        ['数据库驱动', 'PyMySQL', '1.1.x', '免编译'],
        ['WSGI', 'waitress', '3.x+', 'Windows 友好'],
        ['数据库', 'MySQL', '9.4', '库名 campus_market'],
        ['消息队列', 'RabbitMQ', '3.x（可选）', '异步通知 / AI 任务'],
        ['LLM 接入', 'OpenAI 兼容协议', '—', '失败时降级为 mock'],
    ],
    col_widths=[3, 6, 3, 5]
)

add_heading(doc, '2.2.3 前后端分离实现方式', level=3)
add_para(doc, '数据交互：前端通过 wx.request（小程序）/ axios（Web）发起 HTTP 请求，统一前缀 http://127.0.0.1:8000/api/，请求头携带 Authorization: Bearer <access_token>。')
add_para(doc, '统一响应格式：{ "code": 0, "message": "ok", "data": ... }，由 market/response.py 中的自定义 Renderer 实现，业务侧无需自行包装。')
add_para(doc, 'CORS：后端安装 django-cors-headers，默认 CORS_ALLOW_ALL_ORIGINS = True（开发环境），生产环境按域名白名单收敛。')
add_para(doc, '联调主机：miniprogram/app.js 中的 apiBase 需指向电脑局域网 IP（如 http://192.168.x.x:8000），微信开发者工具需勾选「不校验合法域名」。')

add_heading(doc, '2.3 文件组织结构', level=2)
add_heading(doc, '2.3.1 项目根目录树', level=3)
add_code_block(doc, '''综合实训/
├── backend/                # Django 后端（端口 8000）
│   ├── config/             # Django 项目配置
│   ├── market/             # 核心业务 App（12 个模型 + 8 个 view 子模块）
│   │   ├── models.py       # User / Product / Order / Conversation ...
│   │   ├── views/          # auth / product / order / message / ai / admin
│   │   ├── serializers/    # DRF 序列化器
│   │   ├── services/       # ai_service / llm_client / asr_adapter
│   │   └── migrations/
│   ├── finance_legacy/     # 【已下线】原家庭记账代码
│   ├── scripts/            # 建库 SQL / 种子数据
│   └── requirements.txt
├── frontend-web/           # Vue3 + TS 卖家工作台（端口 3000）
├── frontend-admin/         # Vue3 + JS 平台管理后台（端口 5173）
├── miniprogram/            # 微信小程序（原生）
│   ├── pages/              # home / category / publish / messages / mine ...
│   ├── components/         # product-card / voice-input / credit-badge ...
│   ├── custom-tab-bar/     # 5 tab 自定义导航
│   ├── utils/              # api.js / request.js / voice.js
│   └── app.js / app.json
├── deploy/                 # PowerShell 一键脚本
├── docs/                   # 项目文档
└── README.md''')

add_heading(doc, '2.3.2 核心目录与文件作用', level=3)
add_table(doc,
    ['路径', '作用'],
    [
        ['backend/config/settings.py', 'Django 全局配置（DB / JWT / CORS / INSTALLED_APPS）'],
        ['backend/market/models.py', '12 个 ORM 模型，集中体现业务实体'],
        ['backend/market/views/*.py', '按业务领域拆分的 12 个 view 模块'],
        ['backend/market/serializers/*.py', 'DRF 序列化器'],
        ['backend/market/permissions.py', 'IsOwnerOrReadOnly / IsAdminUser 等对象级权限'],
        ['backend/market/services/ai_service.py', 'LLM 客户端封装（OpenAI 兼容协议 + 降级）'],
        ['miniprogram/utils/request.js', '小程序网络请求封装（带 Token 自动续签）'],
        ['miniprogram/custom-tab-bar/', '5 tab 自定义导航 + 中间凸起发布按钮'],
        ['frontend-web/src/api/*.ts', '卖家台 API 模块（强类型）'],
        ['frontend-admin/src/views/AuditProducts.vue', '商品审核页'],
    ],
    col_widths=[7, 8]
)

add_heading(doc, '2.4 开发与运行环境', level=2)
add_heading(doc, '2.4.1 软件与版本', level=3)
add_table(doc,
    ['软件', '版本', '路径 / 备注'],
    [
        ['操作系统', 'Windows 11 64-bit', '中文版'],
        ['Python', '3.13', 'C:\\Users\\liem\\AppData\\Local\\Programs\\Python\\Python313\\python.exe'],
        ['Node.js', '18 LTS+', 'npm 包管理'],
        ['MySQL', '9.4', 'C:\\Program Files\\MySQL\\MySQL Server 9.4\\bin\\；root / tyb1124'],
        ['微信开发者工具', 'Stable 最新', '导入 miniprogram/'],
        ['RabbitMQ（可选）', '3.x', 'D:\\RJ\\RabbitMQ'],
    ],
    col_widths=[3, 2, 10]
)

add_heading(doc, '2.4.2 服务地址与端口', level=3)
add_table(doc,
    ['服务', '地址', '启动方式'],
    [
        ['后端 API', 'http://127.0.0.1:8000/', 'waitress-serve --host=0.0.0.0 --port=8000 --threads=4 config.wsgi:application'],
        ['健康检查', 'http://127.0.0.1:8000/api/health/', '返回 {"code":0,...} 即正常'],
        ['卖家工作台', 'http://127.0.0.1:3000/', 'cd frontend-web && npm run dev'],
        ['管理后台', 'http://127.0.0.1:5173/', 'cd frontend-admin && npm run dev'],
        ['微信小程序', '微信开发者工具内', 'AppID 选「测试号」'],
    ],
    col_widths=[3, 6, 6]
)

add_heading(doc, '2.4.3 常见环境注意事项', level=3)
add_para(doc, '1. PowerShell 不支持 &&：本项目所有 PS 脚本均使用 ; 串联命令，切勿直接复制 && 写法。')
add_para(doc, '2. 跨域：后端默认开启 corsheaders，Web 前端无需配代理即可联调。')
add_para(doc, '3. 小程序域名校验：开发者工具「详情 → 本地设置 → 不校验合法域名」必须勾选；apiBase 须指向电脑局域网 IP。')
add_para(doc, '4. 数据库连接：.env 中 DB_HOST=127.0.0.1 / DB_PORT=3306 / DB_USER=root / DB_PASSWORD=tyb1124 / DB_NAME=campus_market。')
add_para(doc, '5. MEDIA 路径：开发环境 MEDIA_ROOT=backend/media/，生产环境建议改 Nginx 静态托管。')

add_page_break(doc)

# ============= 第3章 =============
add_heading(doc, '第3章 数据库设计', level=1)

add_heading(doc, '3.1 系统 E-R 图', level=2)
add_code_block(doc, '''                            ┌──────────────┐
                            │  market_user │
                            │   (用户)     │
                            └──────┬───────┘
                                   │ 1
                ┌──────────────────┼──────────────────┐
                │ N                │ N                │ N
        ┌───────▼────────┐  ┌──────▼────────┐  ┌──────▼────────┐
        │ market_product │  │market_favorite│  │market_conversation│
        │   (商品)       │  │   (收藏)     │  │   (会话)      │
        └─┬──────────┬───┘  └──────────────┘  └──────┬──────────┘
          │1        │1                               │1
          │N        │N                               │N
  ┌───────▼──────┐  ┌▼──────────────┐         ┌──────▼─────────┐
  │market_product│ │market_category│         │ market_message │
  │  _image      │  │   (分类)      │         │   (消息)        │
  │   (商品图)   │  └───────────────┘         └─────────────────┘
  └──────────────┘
          │ N
          │ 1
  ┌───────▼────────┐         ┌──────────────┐
  │ market_order   │ N     1 │ market_user  │
  │   (订单)       ├─────────┤   (买家)     │
  └───────┬────────┘         └──────────────┘
          │ 1
          │ 1
  ┌───────▼────────┐
  │ market_review  │
  │   (评价)       │
  └────────────────┘

  ┌────────────────┐         ┌────────────────────┐
  │ market_report  │  N   1  │ market_audit_log   │
  │   (举报)       ├─────────┤   (审计日志)       │
  └────────────────┘         └────────────────────┘

  ┌────────────────────────┐
  │ market_system_setting  │  (系统 KV 配置)
  └────────────────────────┘''')
add_para(doc, '关系总结：User 1:N Product；User 1:N Favorite N:1 Product；User 1:N Conversation；Product 1:N Conversation；Product 1:N ProductImage；Product N:1 Category；Product 1:N Order；Order 1:1 Review；Product 1:N Report；User 1:N AuditLog。')

add_heading(doc, '3.2 主要数据表结构', level=2)
add_heading(doc, '3.2.1 用户表 market_user', level=3)
add_table(doc,
    ['字段', '类型', '约束', '说明'],
    [
        ['id', 'BIGINT', 'PK, AUTO_INCREMENT', '主键'],
        ['username', 'VARCHAR(150)', 'UNIQUE, NOT NULL', '登录名'],
        ['password', 'VARCHAR(128)', 'NOT NULL', 'Django pbkdf2 哈希'],
        ['school', 'VARCHAR(64)', 'DEFAULT ""', '学校名称'],
        ['student_id', 'VARCHAR(32)', 'DEFAULT ""', '学号'],
        ['credit_score', 'INT', 'DEFAULT 80', '信用分 0-100'],
        ['avatar', 'VARCHAR(512)', 'DEFAULT ""', '头像 URL'],
        ['role', 'VARCHAR(16)', 'DEFAULT "user"', 'user / admin'],
        ['is_certified', 'TINYINT(1)', 'DEFAULT 0', '是否校园认证'],
        ['is_active', 'TINYINT(1)', 'DEFAULT 1', '启用状态（封禁 = False）'],
        ['created_at / updated_at', 'DATETIME', 'AUTO', '时间戳'],
        ['索引', 'idx_mkt_user_role_active(role, is_active)', '—', '管理员按角色+启用状态过滤'],
    ],
    col_widths=[3, 4, 4, 5]
)
add_para(doc, '设计理由：继承 AbstractUser 复用 Django 密码哈希、Admin 登录；school / student_id 为校园场景扩展；credit_score 用于交易信任；is_active=False 承担封禁语义。')

add_heading(doc, '3.2.2 分类表 market_category', level=3)
add_table(doc,
    ['字段', '类型', '约束', '说明'],
    [
        ['id', 'BIGINT', 'PK', '主键'],
        ['name', 'VARCHAR(32)', 'NOT NULL', '分类名'],
        ['code', 'VARCHAR(32)', 'UNIQUE, NOT NULL', '程序内引用代码'],
        ['parent_id', 'BIGINT', 'FK→self, SET_NULL', '父分类，一级为 NULL'],
        ['icon', 'VARCHAR(64)', 'DEFAULT ""', 'SVG / Lucide 名称（无 emoji）'],
        ['sort_order', 'INT', 'DEFAULT 0', '排序权重'],
        ['is_active', 'TINYINT(1)', 'DEFAULT 1', '是否启用'],
        ['created_at', 'DATETIME', 'AUTO', '创建时间'],
        ['索引', 'idx_mkt_cat_parent_sort(parent_id, sort_order)', '—', '类目树查询'],
    ],
    col_widths=[3, 4, 4, 5]
)
add_para(doc, '设计理由：自引用支持一级 + 二级；code 唯一约束替代硬编码 id；parent 用 SET_NULL 而非 CASCADE。')

add_heading(doc, '3.2.3 商品表 market_product', level=3)
add_table(doc,
    ['字段', '类型', '约束', '说明'],
    [
        ['id', 'BIGINT', 'PK', '主键'],
        ['seller_id', 'BIGINT', 'FK→user, CASCADE', '卖家'],
        ['category_id', 'BIGINT', 'FK→category, PROTECT', '分类（防误删）'],
        ['title', 'VARCHAR(64)', 'NOT NULL', '标题'],
        ['description', 'TEXT(500)', 'DEFAULT ""', '描述'],
        ['price', 'DECIMAL(10,2)', 'NOT NULL', '售价'],
        ['original_price', 'DECIMAL(10,2)', 'NULL', '原价（打折标签）'],
        ['condition', 'VARCHAR(16)', 'DEFAULT "like_new"', 'new/like_new/good/fair'],
        ['status', 'VARCHAR(16)', 'DEFAULT "pending"', 'draft/pending/on_sale/pending_sold/sold/off_shelf'],
        ['view_count', 'INT', 'DEFAULT 0', '浏览数'],
        ['favorite_count', 'INT', 'DEFAULT 0', '收藏数（冗余）'],
        ['created_at / updated_at', 'DATETIME', 'AUTO', '时间戳'],
        ['索引1', 'idx_mkt_prod_status_ctime(status, -created_at)', '—', '首页瀑布流'],
        ['索引2', 'idx_mkt_prod_cat_status(category_id, status)', '—', '类目页'],
        ['索引3', 'idx_mkt_prod_seller_status(seller_id, status)', '—', '我的发布'],
    ],
    col_widths=[3, 4, 4, 5]
)
add_para(doc, '设计理由：状态机字段统一在 status；view_count / favorite_count 冗余避免每次 COUNT；三个复合索引覆盖三大查询路径。')

add_heading(doc, '3.2.4 其他辅助表（结构要点）', level=3)
add_table(doc,
    ['表', '关键字段 / 约束'],
    [
        ['market_product_image', 'product_id (FK CASCADE), image_url, sort_order'],
        ['market_favorite', 'unique_together=(user, product)，级联删除'],
        ['market_conversation', 'unique_together=(product, buyer)，冗余 last_message / unread_*'],
        ['market_message', 'conversation_id + sender_id，索引 (conversation, created_at)'],
        ['market_order', '5 态状态机；price 为下单时快照；索引 (buyer,status) / (seller,status)'],
        ['market_review', 'OneToOne order；rating 1-5 星；reviewer ≠ reviewee'],
        ['market_report', 'reason / status / action 枚举；索引 (status, -created_at)'],
        ['market_audit_log', '不可变日志；记录 operator / action / target_type / target_id'],
        ['market_system_setting', 'KV 存储，AI 开关、Mock 开关、轮播图 JSON'],
    ],
    col_widths=[5, 11]
)

add_heading(doc, '3.3 表关联关系说明', level=2)
add_para(doc, '1. User ↔ Product（一对多）：外键 Product.seller → User.id，on_delete=CASCADE；用户注销后商品不再展示。')
add_para(doc, '2. User ↔ Favorite ↔ Product（多对多）：显式中间表 market_favorite，unique_together=(user, product)；CASCADE 双删。')
add_para(doc, '3. Product ↔ Category（多对一，保护）：Product.category → Category.id，on_delete=PROTECT；阻止删除仍被商品引用的分类。')
add_para(doc, '4. Product ↔ Order（一对多，弱关联）：Order.product → Product.id，on_delete=SET_NULL；商品下架后订单记录仍保留。')
add_para(doc, '5. Order ↔ Review（一对一）：Review.order → Order.id，OneToOneField；同一订单只能写一条评价。')
add_para(doc, '6. Product ↔ Report（一对多）：Report.product → Product.id，on_delete=CASCADE；商品删除时举报记录一并清理。')
add_para(doc, '7. User ↔ AuditLog（多对一，弱关联）：AuditLog.operator → User.id，on_delete=SET_NULL；管理员离职日志保留。')
add_para(doc, '为什么这样设计：状态机不丢历史（订单 SET_NULL 保留交易证据；审计日志 SET_NULL 保留操作记录）；保护核心实体不误删（分类 PROTECT、订单中的买卖家 PROTECT）；冗余字段提升性能（view_count / favorite_count / last_message / school 避免频繁 JOIN 与 COUNT）。', bold=False)

add_page_break(doc)

# ============= 第4章 =============
add_heading(doc, '第4章 系统环境搭建', level=1)

add_heading(doc, '4.1 数据库环境搭建', level=2)
add_para(doc, '步骤 1：安装并启动 MySQL 9.4', bold=True)
add_para(doc, '路径：C:\\Program Files\\MySQL\\MySQL Server 9.4\\bin\\。启动 MySQL 服务，验证：mysql -u root -ptyb1124 -e "SELECT VERSION();"')
add_para(doc, '步骤 2：创建数据库', bold=True)
add_code_block(doc, '"C:\\Program Files\\MySQL\\MySQL Server 9.4\\bin\\mysql.exe" -u root -ptyb1124 < backend\\scripts\\create_mysql_db.sql')

add_para(doc, '关键 SQL（backend/scripts/create_mysql_db.sql）：')
add_code_block(doc, 'CREATE DATABASE IF NOT EXISTS campus_market\n    DEFAULT CHARACTER SET utf8mb4\n    DEFAULT COLLATE utf8mb4_unicode_ci;')

add_para(doc, '步骤 3：执行迁移')
add_code_block(doc, 'cd backend\nC:\\Users\\liem\\AppData\\Local\\Programs\\Python\\Python313\\python.exe manage.py makemigrations market\nC:\\Users\\liem\\AppData\\Local\\Programs\\Python\\Python313\\python.exe manage.py migrate')

add_para(doc, '步骤 4：初始化种子数据')
add_code_block(doc, 'C:\\Users\\liem\\AppData\\Local\\Programs\\Python\\Python313\\python.exe scripts\\init_data_market.py\nC:\\Users\\liem\\AppData\\Local\\Programs\\Python\\Python313\\python.exe scripts\\init_admin.py\nC:\\Users\\liem\\AppData\\Local\\Programs\\Python\\Python313\\python.exe scripts\\init_keywords.py')

add_para(doc, '种子数据包含：管理员 admin/admin123、卖家 zhangsan/123456、lisi/123456、买家 wangwu/123456；46 件带图商品覆盖 5 大分类。')

add_para(doc, '步骤 5：验证')
add_code_block(doc, '"C:\\Program Files\\MySQL\\MySQL Server 9.4\\bin\\mysql.exe" -u root -ptyb1124 -e "USE campus_market; SHOW TABLES;"')

add_heading(doc, '4.2 后端环境搭建', level=2)
add_para(doc, '步骤 1：创建虚拟环境并安装依赖')
add_code_block(doc, 'cd backend\nC:\\Users\\liem\\AppData\\Local\\Programs\\Python\\Python313\\python.exe -m venv .venv\n.\\.venv\\Scripts\\Activate.ps1\npip install -r requirements.txt')

add_para(doc, 'requirements.txt 关键依赖（节选）：')
add_code_block(doc, 'Django>=5.0,<6.0\ndjangorestframework>=13.0\ndjangorestframework-simplejwt>=5.3\ndjango-cors-headers>=4.0\nPyMySQL>=1.1.0\nPillow>=10.0\nwaitress>=3.0\npython-dotenv>=1.0')

add_para(doc, '步骤 2：配置 .env')
add_code_block(doc, '# backend/.env\nSECRET_KEY=django-insecure-replace-in-production\nDEBUG=True\nALLOWED_HOSTS=127.0.0.1,localhost,192.168.*,0.0.0.0\n\nDB_ENGINE=django.db.backends.mysql\nDB_NAME=campus_market\nDB_USER=root\nDB_PASSWORD=tyb1124\nDB_HOST=127.0.0.1\nDB_PORT=3306\n\n# LLM（可选，留空走 mock）\nLLM_API_KEY=\nLLM_BASE_URL=https://api.openai.com/v1\nLLM_MODEL=gpt-4o-mini\n\nMEDIA_ROOT=media\nMEDIA_URL=/media/')

add_para(doc, '步骤 3：启动后端')
add_code_block(doc, '# 方式 A（推荐，waitress 生产级）\n.\\deploy\\start_backend.ps1\n\n# 方式 B（开发，runserver）\nC:\\Users\\liem\\AppData\\Local\\Programs\\Python\\Python313\\python.exe manage.py runserver 0.0.0.0:8000')

add_para(doc, '步骤 4：验证')
add_para(doc, '浏览器访问 http://127.0.0.1:8000/api/health/，返回 {"code":0,"message":"ok","data":{"status":"healthy"}} 即正常。')

add_heading(doc, '4.3 微信小程序环境搭建', level=2)
add_para(doc, '步骤 1：安装微信开发者工具（https://developers.weixin.qq.com/miniprogram/dev/devtools/download.html）。')
add_para(doc, '步骤 2：导入项目 → 项目目录选择 miniprogram/ → AppID 选择「测试号」→ 勾选「不校验合法域名」（必须！）。')
add_para(doc, '步骤 3：配置 miniprogram/app.js')
add_code_block(doc, "App({\n  globalData: {\n    // 关键：将 127.0.0.1 改为电脑局域网 IP\n    apiBase: 'http://192.168.31.100:8000',\n    userInfo: null,\n    token: '',\n  },\n  onLaunch() {\n    const token = wx.getStorageSync('token');\n    if (token) this.globalData.token = token;\n  }\n});")

add_para(doc, '步骤 4：网络请求工具类封装（utils/request.js）')
add_code_block(doc, """// utils/request.js —— 核心逻辑展示
const app = getApp();

function request({ url, method = 'GET', data = null, header = {} }) {
  const baseURL = app.globalData.apiBase;
  return new Promise((resolve, reject) => {
    wx.request({
      url: baseURL + url,
      method,
      data,
      header: {
        'Content-Type': 'application/json',
        'Authorization': app.globalData.token ? `Bearer ${app.globalData.token}` : '',
        ...header,
      },
      success: (res) => {
        if (res.statusCode === 401) {
          wx.removeStorageSync('token');
          wx.navigateTo({ url: '/pages/login/login' });
          return reject(res);
        }
        if (res.data && res.data.code === 0) {
          resolve(res.data.data);
        } else {
          wx.showToast({ title: res.data?.message || '请求失败', icon: 'none' });
          reject(res.data);
        }
      },
      fail: (err) => reject(err),
    });
  });
}

module.exports = { request };""")
add_para(doc, '封装要点：自动拼接 apiBase；自动注入 Authorization: Bearer <token>；401 自动跳转登录；统一处理 {code, message, data} 响应。')

add_heading(doc, '4.4 常见环境问题提示', level=2)
add_table(doc,
    ['问题', '现象', '解决方案'],
    [
        ['Access denied for user root', 'migrate 失败', '检查 .env 中 DB_USER / DB_PASSWORD'],
        ['Unknown database campus_market', 'migrate 失败', '先执行 create_mysql_db.sql'],
        ['No module named pymysql', '启动报错', '重新 pip install -r requirements.txt'],
        ['小程序 request:fail', '预览白屏', '勾选「不校验合法域名」；apiBase 改局域网 IP'],
        ['跨域 CORS', 'Web 端 403', '后端已开启 corsheaders；检查 CORS_ALLOW_ALL_ORIGINS'],
        ['8000 端口占用', '启动失败', 'netstat -ano 查 PID 后结束，或换 --port=8001'],
        ['AI 一键发布返回 mock', '始终降级', '检查 .env 中 LLM_API_KEY 是否配置'],
        ['商品图片 404', '详情页破图', '确认 MEDIA_ROOT=backend/media/ 且 urlpatterns 末尾 static() 挂载'],
        ['PowerShell && 报错', '一键脚本中断', '本项目 PS 脚本均使用 ; 串联，请勿替换为 &&'],
        ['登录提示密码错误', '首次启动', '后端 shell 重置 is_active=True 并 set_password'],
    ],
    col_widths=[4, 3, 8]
)

add_page_break(doc)

# ============= 第5章 =============
add_heading(doc, '第5章 核心模块设计与实现', level=1)
add_para(doc, '本章按「功能说明 → 后端实现 → 小程序实现 → 运行效果」的结构组织，每个模块只展示核心逻辑代码。运行效果截图请参考 docs/ 中的项目演示视频与答辩 PPT。', italic=True)

add_heading(doc, '5.1 用户与认证模块（必做）', level=2)
add_heading(doc, '5.1.1 功能说明', level=3)
add_para(doc, '实现「注册 → 登录 → 资料维护 → 信用分展示」完整链路，核心是基于 SimpleJWT 的 Access + Refresh 双 Token 鉴权。前端通过 Authorization: Bearer <access> 访问受保护接口，401 时自动调用 /api/auth/refresh/ 续签。')

add_heading(doc, '5.1.2 后端实现', level=3)
add_para(doc, '模型（backend/market/models.py，节选 User）：', bold=True)
add_code_block(doc, """class User(AbstractUser):
    school = models.CharField('学校', max_length=64, blank=True, default='')
    student_id = models.CharField('学号', max_length=32, blank=True, default='')
    credit_score = models.IntegerField('信用分', default=80)
    avatar = models.URLField('头像URL', max_length=512, blank=True, default='')
    bio = models.CharField('个人简介', max_length=128, blank=True, default='')
    role = models.CharField('角色', max_length=16,
                            choices=(('user','普通用户'),('admin','管理员')),
                            default='user')
    is_certified = models.BooleanField('是否校园认证', default=False)
    created_at = models.DateTimeField('创建时间', auto_now_add=True)
    updated_at = models.DateTimeField('更新时间', auto_now=True)

    class Meta:
        db_table = 'market_user'
        indexes = [models.Index(fields=['role', 'is_active'], name='idx_mkt_user_role_active')]""")

add_para(doc, '视图（backend/market/views/auth_views.py，节选登录）：', bold=True)
add_code_block(doc, """from rest_framework_simplejwt.tokens import RefreshToken

class LoginView(APIView):
    def post(self, request):
        username = request.data.get('username')
        password = request.data.get('password')
        user = authenticate(username=username, password=password)
        if not user or not user.is_active:
            return api_response(code=1001, message='用户名或密码错误')
        refresh = RefreshToken.for_user(user)
        return api_response(data={
            'access': str(refresh.access_token),
            'refresh': str(refresh),
            'user': UserSerializer(user).data,
        })""")

add_para(doc, 'JWT 鉴权原理（重点）：Access Token 有效期 60 分钟，存于前端 storage；Refresh Token 有效期 7 天，仅在 401 时静默调用 /api/auth/refresh/；SimpleJWT 默认 HS256 + SECRET_KEY 签名；管理员接口在 permission_classes 中追加 IsAdminUser。')

add_para(doc, '路由：')
add_code_block(doc, "path('auth/register/', auth_views.RegisterView.as_view(), name='auth-register'),\npath('auth/login/',    auth_views.LoginView.as_view(),    name='auth-login'),\npath('auth/refresh/',  compat_views.RefreshTokenView.as_view(), name='auth-refresh'),")

add_heading(doc, '5.1.3 小程序实现', level=3)
add_para(doc, 'pages/login/login.js 核心逻辑：')
add_code_block(doc, """const { request } = require('../../utils/request.js');

Page({
  data: { username: '', password: '' },
  onLogin() {
    request({
      url: '/api/auth/login/',
      method: 'POST',
      data: { username: this.data.username, password: this.data.password }
    }).then(res => {
      wx.setStorageSync('token', res.access);
      wx.setStorageSync('refresh', res.refresh);
      getApp().globalData.token = res.access;
      wx.switchTab({ url: '/pages/index/index' });
    });
  }
});""")

add_para(doc, 'utils/request.js 401 拦截 + Token 续签：')
add_code_block(doc, """// 续签逻辑
function refreshToken() {
  return wx.request({
    url: baseURL + '/api/auth/refresh/',
    method: 'POST',
    data: { refresh: wx.getStorageSync('refresh') },
    success: (res) => {
      if (res.data?.code === 0) {
        wx.setStorageSync('token', res.data.data.access);
        getApp().globalData.token = res.data.data.access;
        return res.data.data.access;
      }
    }
  });
}""")

add_heading(doc, '5.1.4 运行效果', level=3)
add_para(doc, '登录页：用户名 + 密码 + 「校园注册」入口；我的页：头像、学校、学号、信用分徽章（彩色 + 数字滚动）；个人资料：可编辑学校 / 学号 / 简介，上传头像；校园认证：上传学生证 → 管理员审核 → 徽章高亮。')

add_heading(doc, '5.2 商品管理模块（必做）', level=2)
add_heading(doc, '5.2.1 功能说明', level=3)
add_para(doc, '支持发布、编辑、上下架、搜索、类目筛选、收藏、浏览数；状态机贯穿商品全生命周期。')

add_heading(doc, '5.2.2 后端实现', level=3)
add_para(doc, '序列化器（serializers/product_serializers.py，节选）：')
add_code_block(doc, """class ProductListSerializer(serializers.ModelSerializer):
    cover = serializers.SerializerMethodField()
    seller_name = serializers.CharField(source='seller.username', read_only=True)
    category_name = serializers.CharField(source='category.name', read_only=True)

    class Meta:
        model = Product
        fields = ['id', 'title', 'price', 'original_price', 'condition',
                  'status', 'cover', 'seller_name', 'category_name',
                  'view_count', 'favorite_count', 'school', 'created_at']

    def get_cover(self, obj):
        img = obj.images.first()
        return img.image_url if img else ''""")

add_para(doc, '视图（views/product_views.py，节选列表 + 创建）：')
add_code_block(doc, """class ProductListCreateView(generics.ListCreateAPIView):
    serializer_class = ProductSerializer
    permission_classes = [IsAuthenticatedOrReadOnly]
    pagination_class = StandardPagination

    def get_queryset(self):
        qs = Product.objects.select_related('seller', 'category')\\
                            .prefetch_related('images').filter(status='on_sale')
        category = self.request.query_params.get('category')
        if category: qs = qs.filter(category__code=category)
        keyword = self.request.query_params.get('q')
        if keyword: qs = qs.filter(Q(title__icontains=keyword) | Q(description__icontains=keyword))
        return qs.order_by('-created_at')

    def perform_create(self, serializer):
        serializer.save(
            seller=self.request.user,
            school=self.request.user.school,
            status='pending',  # 进入待审核
        )""")

add_para(doc, '对象级权限（views/product_views.py，更新/删除）：')
add_code_block(doc, """class ProductDetailView(generics.RetrieveUpdateDestroyAPIView):
    serializer_class = ProductSerializer
    permission_classes = [IsAuthenticatedOrReadOnly, IsOwnerOrReadOnly]
    queryset = Product.objects.all()

    def update(self, request, *args, **kwargs):
        instance = self.get_object()
        if instance.seller != request.user and request.user.role != 'admin':
            return api_response(code=1003, message='无权操作该商品')
        return super().update(request, *args, **kwargs)""")

add_para(doc, '对象级权限实现要点：通过 IsOwnerOrReadOnly 自定义 DRF Permission 类（market/permissions.py），在 has_object_permission 中判断 obj.seller == request.user，未登录用户 / 其他人一律返回 403。')

add_heading(doc, '5.2.3 小程序实现', level=3)
add_para(doc, 'pages/publish/publish.js 核心（AI 一键发布）：')
add_code_block(doc, """Page({
  data: { form: { title:'', description:'', price:'', category:'', images:[] } },

  onAIPublish() {
    if (!this.data.form.images.length) {
      return wx.showToast({ title:'请先上传商品图', icon:'none' });
    }
    wx.showLoading({ title:'AI 正在识别...' });
    request({
      url: '/api/ai/publish-assist/',
      method: 'POST',
      data: { image_url: this.data.form.images[0] }
    }).then(ai => {
      this.setData({
        'form.title': ai.title,
        'form.description': ai.description,
        'form.price': ai.price,
        'form.category': ai.category,
        'aiConfidence': ai.confidence
      });
      wx.hideLoading();
    });
  },

  onSubmit() {
    request({
      url: '/api/products/',
      method: 'POST',
      data: this.data.form
    }).then(() => {
      wx.showToast({ title:'已提交审核' });
      wx.switchTab({ url:'/pages/mine/mine' });
    });
  }
});""")

add_para(doc, 'pages/index/index.js 首页瀑布流：')
add_code_block(doc, """Page({
  data: { products: [], page: 1, hasMore: true },
  onLoad() { this.loadList(); },
  onReachBottom() {
    if (this.data.hasMore) {
      this.setData({ page: this.data.page + 1 });
      this.loadList();
    }
  },
  loadList() {
    request({
      url: `/api/products/?page=${this.data.page}&category=${this.data.activeCategory}`,
    }).then(res => {
      this.setData({
        products: this.data.page === 1 ? res.results : [...this.data.products, ...res.results],
        hasMore: !!res.next,
      });
    });
  }
});""")

add_heading(doc, '5.2.4 运行效果', level=3)
add_para(doc, '首页：双列瀑布流 + 类目快捷入口 + 搜索框 + 轮播公告；分类页：一级 + 二级 tab；发布页：图片上传 + AI 一键按钮 + 手动表单；详情页：图片轮播 + 卖家信息 + 「私聊」「收藏」「想要」按钮。')

add_heading(doc, '5.3 私聊与议价模块（必做）', level=2)
add_heading(doc, '5.3.1 功能说明', level=3)
add_para(doc, '基于商品的私聊会话：买家进入商品详情点「私聊」→ 创建或复用 Conversation → 进入聊天页发送文字消息；会话列表按最后消息时间倒序；未读数实时维护。')

add_heading(doc, '5.3.2 后端实现', level=3)
add_para(doc, '视图（views/message_views.py，节选发送消息）：')
add_code_block(doc, """class SendMessageView(APIView):
    permission_classes = [IsAuthenticated]

    def post(self, request):
        conv_id = request.data.get('conversation_id')
        content = request.data.get('content')
        conv = get_object_or_404(Conversation, id=conv_id)

        if request.user not in (conv.buyer, conv.seller):
            return api_response(code=1003, message='无权发送消息')

        with transaction.atomic():
            msg = Message.objects.create(
                conversation=conv, sender=request.user, content=content
            )
            conv.last_message = content[:200]
            conv.last_message_at = msg.created_at
            if request.user == conv.buyer:
                conv.unread_seller = F('unread_seller') + 1
            else:
                conv.unread_buyer = F('unread_buyer') + 1
            conv.save()
        return api_response(data=MessageSerializer(msg).data)""")

add_para(doc, '默认招呼语（首次进入会话时自动发送）：')
add_code_block(doc, """class ConversationDetailView(generics.RetrieveAPIView):
    def retrieve(self, request, *args, **kwargs):
        conv = self.get_object()
        if not conv.messages.exists() and request.user == conv.buyer:
            Message.objects.create(
                conversation=conv, sender=request.user,
                content=f'你好，我对这件商品「{conv.product.title}」感兴趣~'
            )
            conv.last_message = '你好，我对这件商品感兴趣~'
            conv.last_message_at = timezone.now()
            conv.save()
        return api_response(data=ConversationSerializer(conv).data)""")

add_heading(doc, '5.3.3 小程序实现', level=3)
add_para(doc, 'pages/chat-room/chat-room.js 核心：')
add_code_block(doc, """Page({
  data: { messages: [], inputText: '', conversationId: null },

  onLoad(opts) {
    this.setData({ conversationId: opts.id });
    this.loadMessages();
  },

  loadMessages() {
    request({
      url: `/api/conversations/${this.data.conversationId}/messages/`
    }).then(res => this.setData({ messages: res }));
  },

  onSend() {
    if (!this.data.inputText.trim()) return;
    request({
      url: '/api/messages/send/',
      method: 'POST',
      data: { conversation_id: this.data.conversationId, content: this.data.inputText }
    }).then(msg => {
      this.setData({ messages: [...this.data.messages, msg], inputText: '' });
    });
  }
});""")

add_heading(doc, '5.3.4 运行效果', level=3)
add_para(doc, '会话列表：商品缩略图 + 对方昵称 + 最后消息预览 + 未读数红点；聊天页：左右气泡区分买卖方；议价参考：AI 议价视图在消息流中插入「参考同款历史成交价 ¥XX」灰色提示卡。')

add_heading(doc, '5.4 订单与状态机模块（必做）', level=2)
add_heading(doc, '5.4.1 功能说明', level=3)
add_para(doc, '订单是平台核心交易凭证，状态机 requested → confirmed → shipping → completed，任意阶段可 cancelled。price 字段为下单时快照，避免后续商品改价影响订单。')

add_heading(doc, '5.4.2 后端实现', level=3)
add_para(doc, '模型（节选 Order 状态机）：')
add_code_block(doc, """STATUS_CHOICES = (
    ('requested', '已申请'),
    ('confirmed', '已确认'),
    ('shipping',  '待取/待发'),
    ('completed', '已完成'),
    ('cancelled', '已取消'),
)

class Order(models.Model):
    product = models.ForeignKey(Product, on_delete=models.SET_NULL, null=True, related_name='orders')
    buyer   = models.ForeignKey(User, on_delete=models.PROTECT, related_name='buy_orders')
    seller  = models.ForeignKey(User, on_delete=models.PROTECT, related_name='sell_orders')
    status  = models.CharField('状态', max_length=16, choices=STATUS_CHOICES, default='requested')
    shipping_method = models.CharField('交易方式', max_length=16,
                                       choices=(('pickup','校内自取'),('express','快递')),
                                       default='pickup')
    price = models.DecimalField('成交价', max_digits=10, decimal_places=2)  # 快照
    pickup_location = models.CharField('自取地点', max_length=128, blank=True, default='')
    pickup_time = models.DateTimeField('自取时间', null=True, blank=True)
    created_at = models.DateTimeField('创建时间', auto_now_add=True)
    completed_at = models.DateTimeField('完成时间', null=True, blank=True)

    class Meta:
        db_table = 'market_order'
        indexes = [
            models.Index(fields=['buyer', 'status']),
            models.Index(fields=['seller', 'status']),
        ]""")

add_para(doc, '视图（views/order_views.py，节选确认订单）：')
add_code_block(doc, """class ConfirmOrderView(APIView):
    permission_classes = [IsAuthenticated, IsOrderSeller]

    def post(self, request, pk):
        order = get_object_or_404(Order, pk=pk)
        self.check_object_permissions(request, order)
        with transaction.atomic():
            order.status = 'confirmed'
            order.shipping_method = request.data.get('shipping_method', 'pickup')
            order.pickup_location = request.data.get('pickup_location', '')
            order.pickup_time = request.data.get('pickup_time') or None
            order.save()
            if order.product:
                order.product.status = 'pending_sold'
                order.product.save()
        return api_response(data=OrderSerializer(order).data)""")

add_para(doc, '状态机联动：订单状态变化时，商品状态会同步更新（pending_sold → sold），保证两边数据一致；事务保证「订单 + 商品」原子更新。')

add_heading(doc, '5.4.3 小程序实现', level=3)
add_para(doc, 'pages/orders/orders.js 订单列表：')
add_code_block(doc, """Page({
  data: { orders: [], activeTab: 'all' },
  onLoad() { this.loadOrders(); },
  onTabChange(e) {
    this.setData({ activeTab: e.detail.name });
    this.loadOrders();
  },
  loadOrders() {
    const status = this.data.activeTab === 'all' ? '' : `?status=${this.data.activeTab}`;
    request({ url: `/api/orders/${status}` })
      .then(res => this.setData({ orders: res.results || res }));
  }
});""")
add_para(doc, '订单步骤条组件（components/order-step）：4 步（申请 → 确认 → 待取 → 完成），当前态高亮，已完成打勾。')

add_heading(doc, '5.4.4 运行效果', level=3)
add_para(doc, '我的订单（4 tab）：全部 / 进行中 / 已完成 / 已取消；订单详情：步骤条 + 商品卡片 + 双方信息 + 「确认完成」「取消」按钮；完成后弹出评价弹窗（1-5 星 + 留言）。')

add_heading(doc, '5.5 数据统计与看板模块（必做 + 选做）', level=2)
add_heading(doc, '5.5.1 功能说明', level=3)
add_para(doc, '卖家看板：销售趋势（近 7 / 30 天）、类目分布饼图、价格区间柱图；平台看板：用户数 / 商品数 / 订单数 / 今日新增 / 待审核数 / 趋势。基于 Django ORM 聚合（annotate + Count / Sum / TruncDate）实现，ECharts 渲染。')

add_heading(doc, '5.5.2 后端实现', level=3)
add_para(doc, '视图（views/stats_views.py，节选卖家趋势）：')
add_code_block(doc, """from django.db.models import Count, Sum
from django.db.models.functions import TruncDate

class SellerTrendView(APIView):
    permission_classes = [IsAuthenticated]

    def get(self, request):
        days = int(request.query_params.get('days', 7))
        start = timezone.now() - timedelta(days=days)
        qs = Order.objects.filter(
            seller=request.user, created_at__gte=start,
            status__in=['completed', 'confirmed']
        ).annotate(date=TruncDate('created_at')).values('date')\\
         .annotate(count=Count('id'), amount=Sum('price'))\\
         .order_by('date')
        return api_response(data=list(qs))""")

add_para(doc, 'ORM 聚合要点：TruncDate 按天截断 → values(date) 分组 → Count / Sum 聚合 → 返回 list 给前端直接画图，无需二次计算。')

add_para(doc, 'ECharts 集成（frontend-web/src/views/Statistics.vue 核心）：')
add_code_block(doc, """<template>
  <div ref="trendChart" style="width:100%;height:360px"></div>
</template>

<script setup>
import * as echarts from 'echarts';
import { ref, onMounted } from 'vue';
import { getSellerTrend } from '@/api/stats';

const trendChart = ref(null);
onMounted(async () => {
  const chart = echarts.init(trendChart.value);
  const res = await getSellerTrend(7);
  chart.setOption({
    xAxis: { type: 'category', data: res.map(r => r.date) },
    yAxis: { type: 'value' },
    series: [
      { name:'订单数', type:'line', data: res.map(r => r.count) },
      { name:'销售额', type:'bar',  data: res.map(r => r.amount) },
    ],
  });
});
</script>""")

add_heading(doc, '5.5.3 运行效果', level=3)
add_para(doc, '卖家看板：3 张卡片（今日订单 / 本月销售额 / 在售商品）+ 趋势折线 + 类目饼图 + 价格区间柱图；平台 Dashboard：6 张指标卡 + 趋势 + 热门类目 Top5 + 待审核数。')

add_heading(doc, '5.6 管理后台模块（选做）', level=2)
add_heading(doc, '5.6.1 功能说明', level=3)
add_para(doc, '管理员在 Web 后台（端口 5173）完成商品审核、举报处理、用户封禁、信用分调整、分类管理、AI 配置、审计日志。')

add_heading(doc, '5.6.2 后端实现', level=3)
add_para(doc, '商品审核（views/admin_views.py，节选）：')
add_code_block(doc, """class ProductApproveView(APIView):
    permission_classes = [IsAuthenticated, IsAdminUser]

    def post(self, request, pk):
        product = get_object_or_404(Product, pk=pk)
        with transaction.atomic():
            product.status = 'on_sale'
            product.audited_at = timezone.now()
            product.audit_remark = request.data.get('remark', '')
            product.save()
            AuditLog.objects.create(
                operator=request.user, action='approve_product',
                target_type='product', target_id=product.id, remark=product.audit_remark
            )
        return api_response(data=ProductSerializer(product).data)""")

add_para(doc, '举报处理（带审计）：')
add_code_block(doc, """class ReportHandleView(APIView):
    permission_classes = [IsAuthenticated, IsAdminUser]

    def post(self, request, pk):
        report = get_object_or_404(Report, pk=pk)
        action = request.data.get('action')  # warn/remove/ban/reject
        with transaction.atomic():
            report.status = {'warn':'warned','remove':'removed','ban':'banned','reject':'rejected'}[action]
            report.action = action
            report.remark = request.data.get('remark', '')
            report.handler = request.user
            report.handled_at = timezone.now()
            report.save()
            if action == 'remove' and report.product:
                report.product.status = 'off_shelf'
                report.product.save()
            AuditLog.objects.create(
                operator=request.user, action='handle_report',
                target_type='report', target_id=report.id, remark=report.remark
            )
        return api_response()""")

add_para(doc, '对象级权限实现：所有 admin 视图统一 permission_classes = [IsAuthenticated, IsAdminUser]；非管理员访问返回 403。')

add_heading(doc, '5.6.3 前端实现', level=3)
add_para(doc, 'frontend-admin/src/views/AuditProducts.vue 关键片段：')
add_code_block(doc, """<el-table :data="pendingList">
  <el-table-column prop="title" label="标题" />
  <el-table-column prop="seller_name" label="卖家" />
  <el-table-column prop="created_at" label="提交时间" />
  <el-table-column label="操作">
    <template #default="{ row }">
      <el-button type="success" @click="onApprove(row)">通过</el-button>
      <el-button type="danger"  @click="onReject(row)">驳回</el-button>
    </template>
  </el-table-column>
</el-table>

<script setup>
import { approveProduct, rejectProduct } from '@/api';
const onApprove = (row) => {
  ElMessageBox.confirm('确认通过？').then(() =>
    approveProduct(row.id).then(() => { ElMessage.success('已通过'); loadList(); })
  );
};
const onReject = (row) => {
  ElMessageBox.prompt('请输入驳回理由').then(({ value }) =>
    rejectProduct(row.id, { remark: value }).then(() => loadList())
  );
};
</script>""")

add_heading(doc, '5.6.4 运行效果', level=3)
add_para(doc, '商品审核：表格 + 缩略图 + 驳回理由模板（10 条预设）；举报处理：详情抽屉（举报人 / 商品 / 原因）+ 4 按钮（警告 / 下架 / 封禁 / 驳回）；用户管理：列表 + 封禁 / 解封 + 信用分调整（带「+5 / -10」快捷按钮）；审计日志：所有管理员操作留痕。')

add_heading(doc, '5.7 AI 能力模块（选做）', level=2)
add_heading(doc, '5.7.1 功能说明', level=3)
add_para(doc, '封装 8 个 AI 接口（/api/ai/*），无 LLM Key 时自动降级为 mock 并显式标识「AI 推荐」灰色徽章，避免误导用户。')

add_heading(doc, '5.7.2 后端实现', level=3)
add_para(doc, 'LLM 客户端（services/llm_client.py）：')
add_code_block(doc, """import os
import requests

class LLMClient:
    \"\"\"OpenAI 兼容协议 LLM 客户端封装，支持 mock 降级。\"\"\"
    def __init__(self):
        self.api_key = os.getenv('LLM_API_KEY', '')
        self.base_url = os.getenv('LLM_BASE_URL', 'https://api.openai.com/v1')
        self.model = os.getenv('LLM_MODEL', 'gpt-4o-mini')

    def chat(self, system, user, json_mode=False):
        if not self.api_key:
            return self._mock(system, user)
        try:
            headers = {'Authorization': f'Bearer {self.api_key}'}
            payload = {
                'model': self.model,
                'messages': [
                    {'role': 'system', 'content': system},
                    {'role': 'user', 'content': user},
                ],
            }
            if json_mode: payload['response_format'] = {'type': 'json_object'}
            r = requests.post(f'{self.base_url}/chat/completions',
                              json=payload, headers=headers, timeout=30)
            return r.json()
        except Exception as e:
            return self._mock(system, user, error=str(e))

    def _mock(self, system, user, error=None):
        return {
            'choices': [{'message': {'content':
                '{"title":"闲置物品","description":"九成新","price":50,"category":"other","confidence":0.3}'}}],
            '_mock': True,
            '_error': error,
        }""")

add_para(doc, 'AI 一键发布视图：')
add_code_block(doc, """class AiPublishAssistView(APIView):
    permission_classes = [IsAuthenticated]

    def post(self, request):
        image_url = request.data.get('image_url')
        client = LLMClient()
        result = client.chat(
            system='你是校园二手商品识别助手，根据图片URL返回 JSON 字段。',
            user=f'图片：{image_url}\\n请返回 {{"title","description","category","price","confidence"}}',
            json_mode=True,
        )
        content = result['choices'][0]['message']['content']
        data = json.loads(content)
        data['_mock'] = result.get('_mock', False)
        return api_response(data=data)""")

add_para(doc, '降级策略：无 LLM_API_KEY → 走 mock，返回 _mock=True；网络失败/超时 → catch 异常 → 走 mock；前端收到 _mock=True → 输入框标灰「AI 推荐仅供参考」。')

add_heading(doc, '5.7.3 运行效果', level=3)
add_para(doc, '发布页：上传图片后点「AI 一键发布」→ 自动回填标题 / 描述 / 建议价，标「AI 推荐」灰色徽章；私聊页：议价时 AI 提示「同款历史成交价 ¥XX」；管理后台 AI 配置：可视化测试 LLM Key 是否可用。')

add_page_break(doc)

# ============= 第6章 =============
add_heading(doc, '第6章 系统测试与三端联调', level=1)

add_heading(doc, '6.1 功能测试', level=2)
add_para(doc, '由于篇幅限制，本节列出核心测试用例，完整测试用例集见 backend/test_e2e.py 与 backend/test_e2e_all.py。', italic=True)
add_table(doc,
    ['模块', '测试项', '操作步骤', '预期结果', '实际结果'],
    [
        ['鉴权', '用户注册', 'POST /api/auth/register/ 完整字段', '返回 200 + 自动登录', '通过'],
        ['鉴权', '错误密码登录', '错误密码 POST /api/auth/login/', '返回 1001 用户名或密码错误', '通过'],
        ['鉴权', 'Token 续签', '携带过期 access 调用接口', '401 → 自动 refresh → 重试成功', '通过'],
        ['商品', '发布商品', 'POST /api/products/ 完整字段', 'status=pending 待审核', '通过'],
        ['商品', '未登录发布', '不带 token POST', '返回 401', '通过'],
        ['商品', '编辑他人商品', 'A 发布 B 编辑', '返回 1003 无权操作', '通过'],
        ['商品', '类目筛选', 'GET ?category=textbook_uni', '仅返回该类目', '通过'],
        ['商品', '关键词搜索', 'GET ?q=高等数学', '标题 / 描述命中', '通过'],
        ['收藏', '重复收藏', '同一商品收藏 2 次', '第二次返回 200（幂等）', '通过'],
        ['私聊', '创建会话', '买家对商品发起私聊', '自动创建 / 复用会话', '通过'],
        ['私聊', '发送消息', 'POST /api/messages/send/', '写入消息 + 未读 +1', '通过'],
        ['私聊', '默认招呼语', '首次进入会话', '自动写入招呼语', '通过'],
        ['订单', '下单', 'POST /api/orders/', 'status=requested，商品 status=pending_sold', '通过'],
        ['订单', '卖家确认', 'POST /api/orders/<id>/confirm/', 'status=confirmed', '通过'],
        ['订单', '买家取消 requested', 'POST /api/orders/<id>/cancel/', 'status=cancelled，商品回到 on_sale', '通过'],
        ['订单', '越权确认', '买家调用 confirm', '返回 1003', '通过'],
        ['评价', '完成订单后评价', 'POST /api/reviews/', 'rating 1-5，触发信用分 +1 / -1', '通过'],
        ['统计', '卖家趋势', 'GET /api/stats/seller/trend/?days=7', '返回 7 天日期 + 订单数 + 销售额', '通过'],
        ['统计', '类目分布', 'GET /api/stats/seller/category-distribution/', '返回类目 + 数量 + 占比', '通过'],
        ['管理', '待审核列表', 'GET /api/admin/products/audit/', '仅管理员可见', '通过'],
        ['管理', '普通用户访问', 'user 调用 admin 接口', '返回 1003 需要管理员权限', '通过'],
        ['管理', '审计日志', '通过审核 / 驳回', 'AuditLog 写入对应记录', '通过'],
        ['AI', '一键发布（mock）', '无 Key 调用', '返回 mock + _mock=True', '通过'],
        ['AI', '内容审核', 'POST /api/ai/moderate/ 含敏感词', '标记 risky=True', '通过'],
    ],
    col_widths=[2, 3, 4, 4, 2]
)
add_para(doc, '未通过项：早期版本的图片上传未做格式校验，曾出现上传 PDF 失败；已增加 image/* MIME 校验。', italic=True)

add_heading(doc, '6.2 三端联调测试', level=2)
add_heading(doc, '6.2.1 联调过程', level=3)
add_table(doc,
    ['联调路径', '端', '验证点', '结果'],
    [
        ['后端 ↔ 微信小程序', '8000 ↔ 微信 IDE', 'Token 注入、跨域、局域网 IP 联调', '通过'],
        ['后端 ↔ 卖家工作台', '8000 ↔ 3000', '登录、发布、订单、看板 ECharts', '通过'],
        ['后端 ↔ 管理后台', '8000 ↔ 5173', '审核、封禁、调分、审计', '通过'],
        ['小程序 ↔ 卖家台', '微信 ↔ Web', '同一用户两端的「我的发布」数据一致', '通过'],
        ['卖家台 ↔ 管理后台', 'Web ↔ Web', '卖家发布的商品能在审核列表看到', '通过'],
    ],
    col_widths=[4, 3, 6, 2]
)

add_heading(doc, '6.2.2 联调中发现的问题与解决', level=3)
add_table(doc,
    ['#', '问题', '排查过程', '解决方案'],
    [
        ['1', '小程序 request:fail', '开发者工具未勾选「不校验合法域名」', '勾选后正常；apiBase 改局域网 IP'],
        ['2', '卖家台 401 后白屏', 'Axios 未配置 refresh 拦截器', '在 request.ts 中添加 401 拦截 + refresh + 重试'],
        ['3', '商品图在 Web 端破图', '后端 MEDIA_URL 未挂载到 urlpatterns', '在 config/urls.py 末尾添加 static() 挂载'],
        ['4', '管理后台审核后小程序未刷新', '状态机更新后无消息推送', '当前为轮询方案（首页 onShow 时 reload）'],
        ['5', 'AI 一键发布超时', 'LLM 接口 30s 无响应导致前端 loading 卡死', '给 LLMClient.chat 加 timeout=30；前端加 30s 兜底'],
    ],
    col_widths=[1, 4, 4, 6]
)

add_heading(doc, '6.2.3 重点测试结论', level=3)
add_para(doc, '数据一致性：商品状态流转与订单状态联动准确，事务回滚正常；权限控制：普通用户访问 /api/admin/* 全部 1003 拒绝；A 用户编辑 B 用户的商品 1003 拒绝；信用分联动：好评 +1 / 差评 -1 在订单完成 + 评价提交后即时生效，徽章数字与后端一致。')

add_heading(doc, '6.3 测试结果分析', level=2)
add_heading(doc, '6.3.1 整体运行情况', level=3)
add_para(doc, '必做 5 大模块全部测试通过；选做 3 大模块（AI、管理后台、信用分）全部实现并测试通过；三端联调覆盖买家、卖家、管理员三类用户角色全链路场景；端到端测试脚本 backend/test_e2e_all.py 50+ 用例全部 PASS。')

add_heading(doc, '6.3.2 达到预期目标的功能', level=3)
add_para(doc, '用户与 JWT 鉴权 / 商品状态机与多端发布 / 私聊与议价 / 订单状态机与可视化 / 卖家 + 平台双看板 / 商品审核 + 举报处理 + 审计日志 / AI 一键发布 + 降级 / 信用分徽章 + 数字滚动。')

add_heading(doc, '6.3.3 需改进方向', level=3)
add_para(doc, '性能：商品列表未做 Redis 缓存，高并发下需引入；实时性：私聊消息当前为轮询，应升级 WebSocket；稳定性：AI 任务偶发超时，需引入 RabbitMQ 异步队列与重试机制；可观测性：缺少统一的链路追踪与业务监控大盘。')

add_page_break(doc)

# ============= 第7章 =============
add_heading(doc, '第7章 开发问题与解决方案', level=1)
add_para(doc, '本章真实记录本项目开发过程中遇到的 4 个典型问题，均来源于本项目实际代码与调试日志，非网上通用模板。', italic=True)

add_heading(doc, '问题 1：小程序 request:fail 报 url not in domain list', level=2)
add_para(doc, '问题描述：在微信开发者工具中预览小程序，首页加载瀑布流时所有商品请求均失败，控制台报 request:fail url not in domain list，但同样的 URL 在浏览器中可正常访问。')
add_para(doc, '原因分析：微信小程序要求所有 wx.request 的目标域名必须经过 ICP 备案并在公众平台配置；本项目后端运行在 http://192.168.31.100:8000，属于局域网 IP，既未备案也无法配置；解决方案只有两条路：①申请正式域名并配置；②开发期勾选「不校验合法域名」。')
add_para(doc, '解决方法：', bold=True)
add_para(doc, '1. 微信开发者工具 → 详情 → 本地设置 → 勾选「不校验合法域名、web-view、TLS 版本及 HTTPS 证书」；')
add_para(doc, '2. 在 miniprogram/project.config.json 中固化配置：')
add_code_block(doc, '{\n  "setting": {\n    "urlCheck": false\n  }\n}')
add_para(doc, '3. 真实发布前必须改为正式 HTTPS 域名并完成备案。')
add_para(doc, '排查过程：第一时间用 curl 验证后端能正常返回 → 排除后端问题；搜索官方文档确认是「合法域名」校验 → 定位到「本地设置」。')

add_heading(doc, '问题 2：卖家台 Axios 401 后页面白屏', level=2)
add_para(doc, '问题描述：卖家台登录后访问受保护接口（如 /api/users/me/），Access Token 过期后页面静默白屏，无任何提示。')
add_para(doc, '原因分析：frontend-web/src/utils/request.ts 中仅做了 error => Promise.reject(error)；业务代码在 await api.getMe() 处抛错后没有 catch，Vue 组件无 fallback UI；同时缺少 refresh 拦截器，导致 401 后用户必须手动重新登录。')
add_para(doc, '解决方法：', bold=True)
add_para(doc, '1. 在 request.ts 中添加 401 拦截 + 自动 refresh：')
add_code_block(doc, """axios.interceptors.response.use(
  res => res,
  async error => {
    if (error.response?.status === 401 && !error.config._retry) {
      error.config._retry = true;
      const refresh = localStorage.getItem('refresh');
      const r = await axios.post('/api/auth/refresh/', { refresh });
      localStorage.setItem('token', r.data.data.access);
      error.config.headers.Authorization = `Bearer ${r.data.data.access}`;
      return axios.request(error.config);
    }
    return Promise.reject(error);
  }
);""")
add_para(doc, '2. 业务侧统一加 try/catch + 全局错误提示 ElMessage.error；')
add_para(doc, '3. 路由守卫检测到连续 401 时跳转登录页。')
add_para(doc, '排查过程：复现用后端脚本生成一个 1 分钟后过期的 access，登录后等到过期再访问；观察 Network 面板：401 → 业务代码 throw → 页面无反应；锁定问题在「无 refresh 拦截 + 无 UI 兜底」。')

add_heading(doc, '问题 3：商品图在 Web 端显示 404', level=2)
add_para(doc, '问题描述：卖家台商品详情页 <img :src="product.cover"> 在浏览器开发者工具中显示 404，但后端 media/products/ 目录下文件确实存在。')
add_para(doc, '原因分析：Django 在开发模式下只有 runserver 才自动托管 MEDIA_ROOT；本项目生产环境用 waitress-serve 启动 WSGI，Django 的 urlpatterns 不会自动添加 static() 路由；因此 http://127.0.0.1:8000/media/products/p1_1.jpg 返回 404。')
add_para(doc, '解决方法：', bold=True)
add_para(doc, '在 backend/config/urls.py 末尾添加：')
add_code_block(doc, """from django.conf import settings
from django.conf.urls.static import static

urlpatterns = [
    # ... 业务路由
]

# 仅在 DEBUG 模式下托管 MEDIA（生产应改 Nginx）
if settings.DEBUG:
    urlpatterns += static(settings.MEDIA_URL, document_root=settings.MEDIA_ROOT)""")
add_para(doc, '生产环境推荐改用 Nginx：')
add_code_block(doc, """location /media/ {
    alias D:/xxx/backend/media/;
    expires 7d;
}""")
add_para(doc, '排查过程：ls backend/media/products/ 确认文件存在；浏览器直接访问 /media/products/p1_1.jpg 复现 404；对比 runserver 启动后正常，差异在 WSGI 服务器 → 定位到 media 托管缺失。')

add_heading(doc, '问题 4：AI 一键发布偶发超时导致前端 loading 卡死', level=2)
add_para(doc, '问题描述：发布商品时点击「AI 一键发布」后偶发 30 秒以上无响应，前端 wx.showLoading 一直转圈，用户必须重启小程序。')
add_para(doc, '原因分析：services/llm_client.py 中 requests.post(...) 没有 timeout 参数，默认会无限等待；LLM 服务商偶发网络抖动时，长时间无响应；前端 await 卡住，超时后无任何兜底提示。')
add_para(doc, '解决方法：', bold=True)
add_para(doc, '1. 后端：给 LLM 请求加 30s 超时 + 异常降级到 mock：')
add_code_block(doc, 'r = requests.post(url, json=payload, headers=headers, timeout=30)')
add_para(doc, '2. 后端：在外层 views.ai_views 捕获所有异常，返回降级响应：')
add_code_block(doc, """try:
    result = client.chat(...)
except Exception as e:
    logger.exception('AI service error')
    return api_response(code=0, data=MOCK_RESULT)""")
add_para(doc, '3. 前端：用 Promise.race 强制 30s 超时：')
add_code_block(doc, """const aiPromise = request({ url: '/api/ai/publish-assist/', ... });
const timeout = new Promise((_, reject) =>
  setTimeout(() => reject(new Error('AI_TIMEOUT')), 30000));
Promise.race([aiPromise, timeout])
  .then(...)
  .catch(err => {
    if (err.message === 'AI_TIMEOUT') {
      wx.showToast({ title: 'AI 较慢，请手动填写', icon: 'none' });
    }
  });""")
add_para(doc, '排查过程：后端日志显示 LLM 接口在 20-60s 之间不定时返回（无 timeout）；复现时用 time curl 命令观察到 45s 才返回；在前端加上 loading 时长监控（>5s 提示「AI 识别中，请稍候」），超过 30s 自动取消并兜底。')

add_page_break(doc)

# ============= 第8章 =============
add_heading(doc, '第8章 项目总结', level=1)

add_heading(doc, '8.1 项目完成情况', level=2)
add_para(doc, '已完成功能（覆盖模板全部必做 + 选做）：', bold=True)
add_table(doc,
    ['模块', '状态', '说明'],
    [
        ['用户与 JWT 鉴权', '完成', '注册 / 登录 / 续签 / 校园认证'],
        ['商品管理 + 状态机', '完成', '6 态流转，3 个复合索引'],
        ['私聊与议价', '完成', '会话 / 消息 / 默认招呼语'],
        ['订单状态机', '完成', '5 态 + 自取/快递 + 步骤条'],
        ['数据统计', '完成', '卖家看板 + 平台看板（ECharts）'],
        ['评价与信用分', '完成', '双向评价 + 徽章 + 数字滚动'],
        ['Web 平台管理后台', '完成', '审核 / 举报 / 封禁 / 审计'],
        ['AI 一键发布', '完成', '含 mock 降级'],
        ['三端联调', '完成', '小程序 + Web 卖家台 + Web 后台'],
        ['文档体系', '完成', 'README / QUICKSTART / 设计 Token / Spec'],
    ],
    col_widths=[6, 2, 7]
)
add_para(doc, '未完成 / 简化部分：', bold=True)
add_para(doc, '1. 真实支付集成：仅支持「校内自取 / 快递」线下交易，未接入微信支付担保；')
add_para(doc, '2. WebSocket 私聊：当前私聊为短轮询，未升级 WebSocket；')
add_para(doc, '3. Redis 缓存：商品列表未做缓存，高并发场景未优化；')
add_para(doc, '4. 国际化：仅中文 UI，未做 i18n；')
add_para(doc, '5. 生产部署：仅提供 waitress + 本地 MySQL 方案，未做 Nginx + Gunicorn + RDS 的完整生产化部署。')
add_para(doc, '未完成原因：受课程时间限制，团队将精力集中在「核心链路 + 三端联调演示」，生产化特性作为后续学习方向。', italic=True)

add_heading(doc, '8.2 主要收获', level=2)
add_para(doc, '1. 技术层面', bold=True)
add_para(doc, '深入掌握 Django 5 + DRF 生态：自定义用户、对象级权限、JWT 鉴权、ORM 聚合（TruncDate + Count + Sum）、事务一致性；实践了「一后端 + 多前端」的工程化范式；掌握了 Vue3 Composition API + TypeScript + Pinia + ECharts 的现代前端工程；接触了 LLM 接入、降级策略、Prompt 工程、JSON Mode 等 AI 应用开发基础。')
add_para(doc, '2. 流程层面', bold=True)
add_para(doc, '学会按「需求 → Spec → 任务拆解 → 实现 → 测试 → 联调」的标准流程推进；体会到「业务转型」（家庭记账 → 二手交易）时如何系统化重写 models / serializers / views；掌握了 PowerShell 一键脚本 + 设计 Token 体系，把「环境就绪时间」从 2 小时压缩到 30 秒。')
add_para(doc, '3. 团队协作层面', bold=True)
add_para(doc, '通过「前端 + 后端 + 测试」三角色分工，理解了接口契约（OpenAPI 风格）的重要性；学会了在 git 中按功能分支开发，每个模块独立 PR 评审；体会到代码注释（函数级 docstring）对团队维护的显著价值。')
add_para(doc, '4. 工程素养层面', bold=True)
add_para(doc, '体会到「无 emoji 图标 + 设计 Token」在多端一致性上的威力；学会了用「mock 降级 + 显式标识」保证 AI 不可用时仍能完整演示；形成了「核心链路优先、生产化特性延后」的迭代节奏。')

add_heading(doc, '8.3 存在的不足', level=2)
add_para(doc, '1. 测试覆盖率不足：仅核心模块有 e2e 测试，UI 层缺乏 Playwright / 微信自动化测试，回归风险高；')
add_para(doc, '2. 性能优化欠缺：商品列表、消息列表、统计查询均未做缓存，并发场景下 MySQL 压力大；')
add_para(doc, '3. 可观测性缺失：没有 Sentry / 日志聚合 / 链路追踪，线上问题排查依赖人工复现；')
add_para(doc, '4. 安全加固有限：缺少频率限制、SQL 注入面虽由 ORM 兜底但仍需审计，CORS 仍为 ALLOW_ALL；')
add_para(doc, '5. 生产化欠缺：缺少 CI/CD、Docker 镜像、灰度发布、灾备方案；')
add_para(doc, '6. 业务深度不足：缺少「担保交易」「申诉仲裁」「物流跟踪」「优惠券 / 营销」等更接近真实电商的能力。')

add_heading(doc, '8.4 改进方向', level=2)
add_para(doc, '1. 架构升级', bold=True)
add_para(doc, '引入 Redis 做商品 / 会话 / 热点数据的二级缓存；引入 RabbitMQ（D:\\RJ\\RabbitMQ）做订单状态变更、AI 任务的异步化；私聊升级 WebSocket（dwebsocket / channels）；数据库读写分离 + 慢查询监控。')
add_para(doc, '2. 功能扩展', bold=True)
add_para(doc, '接入微信支付，实现「快递担保交易」；增加「申诉仲裁」流程；引入「校园地图」组件，自取地点可视化；增加「AI 客服 7×24」自动应答 + 人工兜底。')
add_para(doc, '3. 工程化', bold=True)
add_para(doc, '编写 Playwright + 微信自动化端到端测试；引入 GitHub Actions CI，自动化跑 e2e + 覆盖率门禁；Docker Compose 一键起 MySQL + Redis + RabbitMQ + Django + Nginx；接入 Sentry 收集前后端异常。')
add_para(doc, '4. 产品体验', bold=True)
add_para(doc, '增加「暗色模式」+ 「系统主题跟随」；增加「商品对比」「收藏夹分组」「价格提醒」；增加「信用分提升任务」；增加「校园认证」学生证 OCR 识别。')
add_para(doc, '5. 教学沉淀', bold=True)
add_para(doc, '沉淀「课设项目模板」，下一届同学可在此基础上改造；输出 5 分钟答辩演示视频 + 项目答辩 PPT；形成「常见问题 FAQ」文档，沉淀本次遇到的 4 个典型问题及解法。')

# 页脚说明
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
pf = p.paragraph_format
pf.space_before = Pt(40)
run = p.add_run('—— 全文完 ——')
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('项目版本：v2.0（业务转型版）  ·  最后更新：2026-06-13  ·  综合实训项目组')
run.font.size = Pt(9)
run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

# 保存
out_path = r'd:\文件\工作 作业\微信小程序实训\4次课程内容\综合实训\校园二手交易平台-实训报告.docx'
doc.save(out_path)
print(f'OK: {out_path}')
